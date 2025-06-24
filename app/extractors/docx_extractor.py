# app/extractors/docx_extractor.py
from docx import Document
from docx.shared import Inches
import logging
from typing import Dict, List

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file with structure preservation."""
    try:
        doc = Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_texts.append(" | ".join(row_text))
        
        # Combine all text
        all_text = paragraphs + table_texts
        text = "\n".join(all_text)
        
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""

def extract_metadata_from_docx(file_path: str) -> Dict:
    """Extract metadata from DOCX file."""
    metadata = {}
    
    try:
        doc = Document(file_path)
        core_props = doc.core_properties
        
        metadata.update({
            'title': core_props.title or '',
            'author': core_props.author or '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'comments': core_props.comments or '',
            'created': core_props.created.isoformat() if core_props.created else '',
            'modified': core_props.modified.isoformat() if core_props.modified else '',
            'last_modified_by': core_props.last_modified_by or '',
            'revision': core_props.revision or '',
            'category': core_props.category or ''
        })
        
    except Exception as e:
        logger.error(f"DOCX metadata extraction failed: {e}")
    
    return metadata
